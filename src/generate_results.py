"""
Generate the Results section (Section 4) as a Word document,
matching the formatting of the edited methodology document.

Data sources:
  - outputs/combined_comparison.json   → main classifier results (n = 2,332)
  - outputs/hard_subset_comparison.json → adversarial hard-subset (n = 100)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import json, os

# ── Load experimental data ───────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR   = os.path.dirname(SCRIPT_DIR)
IMG_DIR    = os.path.join(ROOT_DIR, "outputs", "figures")
OUT_DIR    = os.path.join(ROOT_DIR, "outputs")

with open(os.path.join(OUT_DIR, "combined_comparison.json")) as f:
    MAIN = json.load(f)

hard_path = os.path.join(OUT_DIR, "hard_subset_comparison.json")
if os.path.exists(hard_path):
    with open(hard_path) as f:
        HARD = json.load(f)
else:
    HARD = None
    print("WARNING: hard_subset_comparison.json not found — skipping §4.5")

# Load RAG evaluation data
rag_path = os.path.join(OUT_DIR, "rag_evaluation.json")
if os.path.exists(rag_path):
    with open(rag_path) as f:
        RAG = json.load(f)
else:
    RAG = None
    print("WARNING: rag_evaluation.json not found — skipping §4.6")

# ── Derived metrics (from combined_comparison.json) ──────────────────────────
NAMES_ORDER = ["XGBoost", "RandomForest", "LightGBM", "LogisticRegression"]
NICE_NAMES  = {"XGBoost": "XGBoost", "RandomForest": "RF",
               "LightGBM": "LightGBM", "LogisticRegression": "LR"}

MARKER_NAMES = [
    ("standardized_structure",    "Standardized Structure",    "x₁"),
    ("predictable_criticism",     "Predictable Criticism",     "x₂"),
    ("excessive_balance",         "Excessive Balance",         "x₃"),
    ("linguistic_homogeneity",    "Linguistic Homogeneity",    "x₄"),
    ("generic_domain_language",   "Generic Domain Language",   "x₅"),
    ("conceptual_feedback",       "Conceptual Feedback",       "x₆"),
    ("absence_personal_signals",  "Absence of Personal Signals","x₇"),
    ("repetition_patterns",       "Repetition Patterns",       "x₈"),
]


def _metrics(d):
    """Compute precision, recall, F1 for the AI-Generated (positive) class
    from the confusion matrix [[TN, FP],[FN, TP]]."""
    cm = d["confusion_matrix"]
    tn, fp = cm[0]
    fn, tp = cm[1]
    n = tn + fp + fn + tp
    acc = d["accuracy"]
    auc = d["auc_roc"]
    prec = tp / (tp + fp) if tp + fp else 0
    rec  = tp / (tp + fn) if tp + fn else 0
    f1   = 2 * prec * rec / (prec + rec) if prec + rec else 0
    fpr  = fp / (fp + tn) if fp + tn else 0
    fnr  = fn / (fn + tp) if fn + tp else 0
    return dict(acc=acc, auc=auc, prec=prec, rec=rec, f1=f1,
                tp=tp, fp=fp, fn=fn, tn=tn, fpr=fpr, fnr=fnr, n=n)


# ── Document setup ───────────────────────────────────────────────────────────
M_NS      = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
BODY_SIZE  = Pt(11)
CAPTION_SIZE = Pt(9)

doc = Document()

section = doc.sections[0]
section.page_width   = Cm(21.0)
section.page_height  = Cm(29.7)
section.left_margin  = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin   = Cm(2.5)
section.bottom_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.space_after  = Pt(0)
style.paragraph_format.space_before = Pt(0)

# ── Helpers ──────────────────────────────────────────────────────────────────

def add_heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = BODY_SIZE

def add_heading2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = BODY_SIZE

def body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = BODY_SIZE
    return p

def body_cont(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = BODY_SIZE
    return p

def body_mixed(segments):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.0
    for text, bold, italic in segments:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = BODY_SIZE
        run.bold   = bold
        run.italic = italic
    return p

def bullet(bold_prefix, text):
    p = doc.add_paragraph(style="List Paragraph")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.name = "Times New Roman"
        rb.font.size = BODY_SIZE
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = BODY_SIZE

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = BODY_SIZE

def figure(image_path, caption_text, width_cm=14):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after  = Pt(2)
    run = p_img.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.line_spacing = 1.0
    p_cap.paragraph_format.space_after  = Pt(4)
    run_cap = p_cap.add_run(caption_text)
    run_cap.font.name = "Times New Roman"
    run_cap.font.size = BODY_SIZE

# ── Three-line table (booktabs) ──────────────────────────────────────────────

def _set_cell_border(cell, **kwargs):
    tc_pr = cell._element.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = tc_pr.makeelement(qn("w:tcBorders"), {})
        tc_pr.append(tc_borders)
    for edge, attrs in kwargs.items():
        elem = tc_borders.makeelement(qn(f"w:{edge}"), {
            qn(f"w:{k}"): str(v) for k, v in attrs.items()
        })
        existing = tc_borders.find(qn(f"w:{edge}"))
        if existing is not None:
            tc_borders.remove(existing)
        tc_borders.append(elem)

_NB = {"val": "nil",    "sz": "0",  "color": "auto",   "space": "0"}
_TB = {"val": "single", "sz": "12", "color": "000000", "space": "0"}
_LB = {"val": "single", "sz": "6",  "color": "000000", "space": "0"}

def add_table(headers, rows):
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = CAPTION_SIZE
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(str(val))
                run.italic = True if val else False
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = CAPTION_SIZE
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            borders = {"start": _NB, "end": _NB}
            if r_idx == 0:
                borders["top"] = _TB;  borders["bottom"] = _LB
            elif r_idx == n_rows - 1:
                borders["top"] = _NB;  borders["bottom"] = _TB
            else:
                borders["top"] = _NB;  borders["bottom"] = _NB
            _set_cell_border(cell, **borders)
    return table


# ═══════════════════════════════════════════════════════════════════════════════
#  Pre-compute all values from JSON so nothing is hardcoded
# ═══════════════════════════════════════════════════════════════════════════════

# Main test-set metrics
M = {name: _metrics(MAIN[name]) for name in NAMES_ORDER}
N_TEST = M["XGBoost"]["n"]  # 2,332

# Hard-subset metrics (optional)
if HARD is not None:
    H = {name: _metrics(HARD[name]) for name in NAMES_ORDER}
    N_HARD = H["XGBoost"]["n"]  # 100
else:
    H = None
    N_HARD = 0

# Feature importance (native, from combined_comparison.json)
def _sorted_importance(d):
    fi = d["feature_importance"]
    total = sum(fi.values())
    ranked = sorted(fi.items(), key=lambda x: x[1], reverse=True)
    return [(k, v, v/total*100) for k, v, in [(k, v) for k, v in ranked]]

XGB_FI = _sorted_importance(MAIN["XGBoost"])


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4: RESULTS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading1("4. Results")

body(
    "This section presents the experimental results of the proposed AI-generated "
    "peer review detection framework. We report the findings in six parts: "
    "dataset overview (§4.1), descriptive analysis of the marker distributions "
    "(§4.2), classifier comparison on the full dataset (§4.3), SHAP-based "
    "explainability analysis (§4.4), adversarial robustness evaluation (§4.5), "
    "and RAG retrieval evaluation (§4.6)."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4.1 Dataset Overview
# ═══════════════════════════════════════════════════════════════════════════════

add_heading2("4.1. Dataset overview")

body(
    "Figure XX provides a visual summary of the dataset composition. "
    "Panel (a) shows the class distribution: 5,772 human-authored reviews and "
    "2,000 AI-generated reviews, yielding a class ratio of approximately 74:26. "
    "Panel (b) disaggregates the data by source, revealing that the human class "
    "is dominated by ICLR 2017 reviews (n = 5,458), supplemented by smaller "
    "contributions from ACL 2017 (n = 275) and CoNLL 2016 (n = 39). The AI class "
    "comprises 1,000 neutral Gen-Review reviews and 1,000 adversarial reviews "
    "(500 from Claude Sonnet 4.6 and 500 from GPT 5.4). Panel (c) shows "
    "the review length distributions: human reviews exhibit a right-skewed "
    "distribution peaking between 100 and 300 words, while AI-generated reviews "
    "concentrate in a narrower range around 500–800 words, reflecting the more "
    "consistent output length of LLM generation."
)

figure(
    os.path.join(IMG_DIR, "fig1_dataset_composition.png"),
    "Figure XX. Dataset composition: (a) class distribution (Human vs. AI-Generated), "
    "(b) source distribution showing the contribution of each data subset, and "
    "(c) review length distributions by class.",
    width_cm=16,
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4.2 Marker Descriptive Analysis
# ═══════════════════════════════════════════════════════════════════════════════

add_heading2("4.2. Descriptive analysis of marker distributions")

body(
    "Before training classifiers, we examined the distributional properties "
    "of the eight marker scores across the human and AI-generated review "
    "classes. Table XX reports the mean and standard deviation of each marker, "
    "together with the mean difference Δ between AI-generated and human reviews."
)

caption("Table XX. Marker descriptive statistics by class (mean ± standard deviation)")

add_table(
    headers=["Marker", "Human (n = 5,772)", "AI (n = 2,000)", "Δ (AI − Human)"],
    rows=[
        ["Standardized Structure (x₁)", "0.205 ± 0.227", "0.457 ± 0.393", "+0.252"],
        ["Predictable Criticism (x₂)",  "0.213 ± 0.146", "0.427 ± 0.249", "+0.215"],
        ["Excessive Balance (x₃)",      "0.268 ± 0.178", "0.512 ± 0.243", "+0.244"],
        ["Linguistic Homogeneity (x₄)", "0.524 ± 0.128", "0.554 ± 0.258", "+0.030"],
        ["Generic Domain Language (x₅)", "0.284 ± 0.148", "0.473 ± 0.323", "+0.189"],
        ["Conceptual Feedback (x₆)",    "0.502 ± 0.208", "0.581 ± 0.266", "+0.079"],
        ["Absence Personal Signals (x₇)", "0.405 ± 0.265", "0.469 ± 0.425", "+0.065"],
        ["Repetition Patterns (x₈)",    "0.167 ± 0.112", "0.398 ± 0.300", "+0.231"],
    ],
)

body(
    "Several observations emerge. First, all eight markers show higher mean "
    "values for AI-generated reviews compared to human reviews, confirming "
    "the directional hypotheses underlying the taxonomy. Second, the magnitude "
    "of separation varies considerably across markers. Standardized Structure "
    "(Δ = +0.252), Excessive Balance (Δ = +0.244), and Repetition Patterns "
    "(Δ = +0.231) exhibit the largest mean differences, suggesting these are "
    "the most discriminative markers at the population level."
)

body(
    "Figure XX presents the per-marker score distributions as overlaid "
    "histograms for human (green) and AI-generated (red) reviews. The "
    "distributions reveal that markers such as Standardized Structure and "
    "Repetition Patterns show clearly separated modes between the two classes, "
    "while Linguistic Homogeneity and Conceptual Feedback exhibit substantial "
    "overlap, indicating that these markers alone provide weaker class "
    "separation. Notably, several AI-marker distributions are bimodal, "
    "reflecting the mixture of neutral and adversarial AI reviews."
)

figure(
    os.path.join(IMG_DIR, "fig2_marker_distributions.png"),
    "Figure XX. Per-marker score distributions for human (green) and AI-generated "
    "(red) reviews, shown as overlaid histograms for each of the eight markers.",
    width_cm=16,
)

body(
    "Figure XX provides a complementary view through box plots, which "
    "summarise the median, interquartile range, and outliers for each "
    "marker by class. The box plots confirm the pattern observed in the "
    "histograms: markers with large mean differences (Standardized Structure, "
    "Repetition Patterns) show non-overlapping interquartile ranges, while "
    "markers with smaller differences (Linguistic Homogeneity, Conceptual "
    "Feedback) exhibit overlapping boxes but still differ in their median "
    "values."
)

figure(
    os.path.join(IMG_DIR, "fig3_marker_boxplot.png"),
    "Figure XX. Box plot comparison of the eight marker scores by class "
    "(Human vs. AI-Generated). Boxes indicate the interquartile range (IQR); "
    "whiskers extend to 1.5× IQR; individual points indicate outliers.",
    width_cm=16,
)

body(
    "Notably, Linguistic Homogeneity (Δ = +0.030) and Absence of Personal "
    "Signals (Δ = +0.065) show relatively small differences in their overall "
    "means. However, these aggregate statistics mask a crucial compositional "
    "effect: the adversarial AI reviews were specifically designed to minimise "
    "these markers, while the neutral Gen-Review reviews exhibit extremely high "
    "scores (e.g., Absence of Personal Signals: 0.894 for neutral AI vs. 0.048 "
    "for adversarial Claude). This bimodal distribution within the AI class "
    "reduces the aggregate mean difference but does not diminish the marker's "
    "discriminative power for non-adversarial detection."
)

body(
    "Table XX disaggregates the AI-generated reviews by source to illustrate "
    "this compositional effect."
)

caption("Table XX. Marker means by AI data source")

add_table(
    headers=["Marker", "Gen-Review\n(neutral, n=1,000)", "Adversarial\nClaude 4.6 (n=500)", "Adversarial\nGPT 5.4 (n=500)"],
    rows=[
        ["Standardized Structure",   "0.849", "0.077", "0.053"],
        ["Predictable Criticism",    "0.665", "0.201", "0.177"],
        ["Excessive Balance",        "0.748", "0.262", "0.292"],
        ["Linguistic Homogeneity",   "0.810", "0.298", "0.299"],
        ["Generic Domain Language",  "0.793", "0.164", "0.143"],
        ["Conceptual Feedback",      "0.840", "0.330", "0.315"],
        ["Absence Personal Signals", "0.894", "0.048", "0.042"],
        ["Repetition Patterns",      "0.696", "0.105", "0.096"],
    ],
)

body(
    "The contrast is stark. Neutral Gen-Review reviews score above 0.65 on "
    "all eight markers, with Absence of Personal Signals (0.894) and "
    "Standardized Structure (0.849) reaching near-ceiling values. In contrast, "
    "both adversarial sources produce marker scores that are comparable to, "
    "or even lower than, human reviews on several markers (e.g., Standardized "
    "Structure: 0.077 for adversarial Claude vs. 0.205 for human). This confirms "
    "that the adversarial prompts successfully induced human-mimicking behaviour, "
    "creating a genuinely challenging detection scenario. The two adversarial "
    "sources (Claude Sonnet 4.6 and GPT 5.4) produce remarkably similar marker "
    "profiles, suggesting that the evasion strategies generalise across LLM "
    "architectures."
)

body(
    "Figure XX presents the inter-marker correlation matrix. The highest "
    "pairwise correlations are observed between Repetition Patterns and "
    "Standardized Structure (r = 0.83) and between Predictable Criticism and "
    "Generic Domain Language (r = 0.83), indicating that reviews with formulaic "
    "structure also tend to exhibit repetitive phrasing and that generic domain "
    "vocabulary co-occurs with predictable criticism patterns. Generic Domain "
    "Language and Conceptual Feedback also show strong correlation (r = 0.81), "
    "suggesting that reviews relying on domain-generic vocabulary tend to "
    "provide more superficial, template-driven feedback. Linguistic Homogeneity "
    "and Absence of Personal Signals (r = 0.78), as well as Excessive Balance "
    "and Generic Domain Language (r = 0.78), form another cluster of "
    "moderately correlated features. The lowest pairwise correlations involve "
    "Excessive Balance and Absence of Personal Signals (r = 0.41) and "
    "Conceptual Feedback and Standardized Structure (r = 0.41), confirming "
    "that these markers capture complementary signal dimensions. The generally "
    "moderate-to-strong positive correlations indicate that AI-generated text "
    "tends to trigger multiple markers simultaneously, while the absence of "
    "near-unity correlations confirms that no two markers are redundant."
)

figure(
    os.path.join(IMG_DIR, "fig11_correlation_heatmap.png"),
    "Figure XX. Pearson correlation matrix of the eight marker scores computed "
    "over the full dataset (N = 7,772). Values are displayed in the lower "
    "triangle to avoid redundancy.",
    width_cm=12,
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4.3 Classifier Performance  (data-driven from combined_comparison.json)
# ═══════════════════════════════════════════════════════════════════════════════

add_heading2("4.3. Classifier comparison")

body(
    f"Table XX presents the performance of the four classifiers on the held-out "
    f"test set (n = {N_TEST:,}). The three ensemble methods, XGBoost, Random Forest (RF), "
    f"and LightGBM, achieve near-identical performance, while the Logistic "
    f"Regression baseline performs substantially lower."
)

caption(f"Table XX. Classifier performance on the held-out test set (n = {N_TEST:,})")

# Build rows from JSON data
perf_rows = []
for name in NAMES_ORDER:
    m = M[name]
    perf_rows.append([
        NICE_NAMES[name],
        f"{m['acc']:.4f}",
        f"{m['auc']:.4f}",
        f"{m['prec']:.4f}",
        f"{m['rec']:.4f}",
        f"{m['f1']:.4f}",
    ])

add_table(
    headers=["Classifier", "Accuracy", "AUC-ROC", "Precision", "Recall", "F₁-Score"],
    rows=perf_rows,
)

# Extract key values for narrative
xgb_m  = M["XGBoost"]
rf_m   = M["RandomForest"]
lgb_m  = M["LightGBM"]
lr_m   = M["LogisticRegression"]

body(
    f"The three tree-based ensemble classifiers achieve an identical accuracy "
    f"of {xgb_m['acc']*100:.2f}% and AUC-ROC values exceeding 0.999. These results "
    f"demonstrate that the eight-marker feature representation provides sufficient "
    f"discriminative power for near-perfect binary classification. The marginal "
    f"differences between XGBoost, RF, and LightGBM are not "
    f"statistically meaningful at this performance level."
)

body(
    "Figure XX provides a visual comparison of three key performance metrics "
    ", Accuracy, AUC-ROC, and F₁-Score, across classifiers through grouped "
    "bar charts. Panels (a) and (b) confirm the near-ceiling performance of "
    f"the tree-based ensembles (all above 0.99), while the Logistic Regression (LR) "
    f"baseline shows a visible gap ({lr_m['acc']:.4f} accuracy and "
    f"{lr_m['auc']:.4f} AUC-ROC). Panel (c) highlights the F₁-Score disparity "
    f"most starkly: all three ensemble methods achieve {xgb_m['f1']:.4f}, whereas "
    f"LR falls to {lr_m['f1']:.4f}, underscoring the precision "
    f"deficit of the linear model."
)

figure(
    os.path.join(IMG_DIR, "fig7_model_comparison.png"),
    "Figure XX. Classifier performance comparison: (a) Accuracy, (b) AUC-ROC, "
    "and (c) F₁-Score for all four classifiers.",
    width_cm=16,
)

body(
    f"LR, serving as a linear baseline, achieves a considerably "
    f"lower accuracy of {lr_m['acc']*100:.2f}% and AUC-ROC of {lr_m['auc']:.3f}. "
    f"While the recall is high ({lr_m['rec']:.3f}), the precision drops to "
    f"{lr_m['prec']:.3f}, resulting in an F₁-Score of {lr_m['f1']:.3f}. "
    f"This suggests that the decision boundary between human and AI reviews in "
    f"the eight-dimensional marker space is non-linear, which the ensemble "
    f"methods capture effectively through their tree-based partitioning."
)

body(
    f"Figure XX presents the Receiver Operating Characteristic (ROC) curves "
    f"for all four classifiers. The three ensemble methods, XGBoost "
    f"(AUC = {xgb_m['auc']:.4f}), RF (AUC = {rf_m['auc']:.4f}), "
    f"and LightGBM (AUC = {lgb_m['auc']:.4f}), produce near-rectangular "
    f"curves that hug the top-left corner, indicating near-perfect "
    f"discrimination across all classification thresholds. The three curves "
    f"are virtually indistinguishable, reflecting the equivalent performance "
    f"of the ensemble methods. LR (AUC = {lr_m['auc']:.4f}) "
    f"exhibits a visibly different trajectory: while it achieves high true "
    f"positive rates, it does so at the cost of substantially higher false "
    f"positive rates, with the curve deviating noticeably from the top-left "
    f"corner. The dashed diagonal represents the random baseline (AUC = 0.5), "
    f"against which all classifiers show a dramatic separation."
)

figure(
    os.path.join(IMG_DIR, "fig4_roc_curves.png"),
    "Figure XX. Receiver Operating Characteristic (ROC) curves for all four "
    "classifiers, with AUC scores shown in the legend. The dashed diagonal "
    "represents random chance (AUC = 0.5).",
    width_cm=12,
)

body_mixed([
    ("Confusion matrix analysis. ", True, False),
    (f"Table XX reports the confusion matrices for all four classifiers. The "
     f"ensemble methods produce remarkably few errors: XGBoost misclassifies "
     f"only {xgb_m['fp']} human reviews as AI (false positives) and "
     f"{xgb_m['fn']} AI reviews as human (false negatives), out of "
     f"{N_TEST:,} test instances. RF and LightGBM each produce "
     f"{rf_m['fp']} false positives and {rf_m['fn']} false negatives.", False, False),
])

caption("Table XX. Confusion matrices on the held-out test set")

# Build confusion matrix table from JSON data
cm_rows = []
for name in NAMES_ORDER:
    m = M[name]
    cm_rows.append([
        NICE_NAMES[name],
        str(m['tp']),
        str(m['fp']),
        str(m['fn']),
        f"{m['tn']:,}",
        f"{m['fpr']*100:.2f}%",
        f"{m['fnr']*100:.2f}%",
    ])

add_table(
    headers=["Classifier", "TP", "FP", "FN", "TN", "FPR", "FNR"],
    rows=cm_rows,
)

figure(
    os.path.join(IMG_DIR, "fig6_confusion_matrices.png"),
    "Figure XX. Confusion matrices for all four classifiers on the held-out "
    "test set. Cell intensities are proportional to the count.",
    width_cm=16,
)

body(
    f"The false positive rate (FPR) of the ensemble methods is below "
    f"{max(xgb_m['fpr'], rf_m['fpr'], lgb_m['fpr'])*100:.2f}%, meaning that "
    f"fewer than 1 in 400 legitimate human reviews would be incorrectly flagged "
    f"as AI-generated. This is a critical practical property: a high false "
    f"positive rate would undermine trust in the detection system by wrongly "
    f"accusing human reviewers. The false negative rate (FNR) of approximately "
    f"{xgb_m['fnr']*100:.1f}% indicates that a small number of AI-generated "
    f"reviews may evade detection, which is expected given the adversarial subset."
)

body(
    f"In contrast, LR produces {lr_m['fp']} false positives "
    f"(FPR = {lr_m['fpr']*100:.2f}%), which would be unacceptable in a "
    f"production deployment where over one in ten legitimate reviews would be "
    f"flagged. This further underscores the need for non-linear classifiers in "
    f"this detection task."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4.4 SHAP Analysis
# ═══════════════════════════════════════════════════════════════════════════════

add_heading2("4.4. Feature importance and SHAP analysis")

# ── 4.4.1 Native feature importance (from combined_comparison.json) ──────────

body(
    "To understand which markers drive classifier decisions, we examine "
    "two complementary perspectives: (i) native model feature importance and "
    "(ii) SHAP-based explainability analysis. Table XX reports the native "
    "feature importance for XGBoost, which reflects the aggregate gain "
    "contributed by each feature across all trees."
)

caption("Table XX. Native feature importance, XGBoost (gain-based)")

# Build native importance table from JSON
xgb_fi = MAIN["XGBoost"]["feature_importance"]
xgb_fi_total = sum(xgb_fi.values())
xgb_fi_sorted = sorted(xgb_fi.items(), key=lambda x: x[1], reverse=True)

MARKER_NICE = {k: f"{nice} ({sym})" for k, nice, sym in MARKER_NAMES}

fi_rows = []
for rank, (key, val) in enumerate(xgb_fi_sorted, 1):
    pct = val / xgb_fi_total * 100
    fi_rows.append([
        str(rank),
        MARKER_NICE[key],
        f"{val:.4f}",
        f"{pct:.1f}%",
    ])

add_table(
    headers=["Rank", "Marker", "Importance", "Relative contribution"],
    rows=fi_rows,
)

# Extract top markers for narrative
top1_key, top1_val = xgb_fi_sorted[0]
top2_key, top2_val = xgb_fi_sorted[1]
top1_pct = top1_val / xgb_fi_total * 100
top2_pct = top2_val / xgb_fi_total * 100

top1_name = dict(MARKER_NAMES[:], )[top1_key] if False else {k: nice for k, nice, _ in MARKER_NAMES}[top1_key]
top2_name = {k: nice for k, nice, _ in MARKER_NAMES}[top2_key]

body(
    f"{top1_name} emerges as the most influential marker in the native "
    f"XGBoost importance, contributing {top1_pct:.1f}% of the total gain. "
    f"This is followed by {top2_name} ({top2_pct:.1f}%). "
    f"Together, these two markers account for over {top1_pct + top2_pct:.0f}% "
    f"of the model's total feature importance, indicating that the XGBoost "
    f"ensemble heavily relies on these features to partition the decision space."
)

body(
    "The native importance ranking, however, captures only aggregate gain, "
    "a measure of how frequently and effectively a feature is used for splitting "
    "across all trees. It does not account for feature interactions or the "
    "direction of effect. To address these limitations, we complement the native "
    "importance analysis with SHAP (SHapley Additive exPlanations), which "
    "provides instance-level, theoretically grounded attribution scores."
)

# ── 4.4.2 SHAP global importance ────────────────────────────────────────────

body(
    "Figure XX presents the global SHAP feature importance (mean |SHAP value|) "
    "side-by-side for all three tree-based classifiers, colour-coded by marker "
    "category (Structural, Argumentative, Linguistic, Behavioural). "
    "A notable distinction emerges between the SHAP and native importance "
    "rankings for XGBoost: SHAP assigns the highest importance to Standardized "
    "Structure (mean |SHAP| = 3.187), whereas the native gain-based importance "
    "ranks Absence of Personal Signals first. This discrepancy arises because "
    "SHAP measures the average marginal contribution of each feature to "
    "individual predictions, accounting for feature correlations and "
    "interactions, while native importance measures aggregate split gain. "
    "The high native importance of Absence of Personal Signals reflects its "
    "frequent use as a splitting variable, but its redundancy with other markers "
    "(r = 0.78 with Linguistic Homogeneity) means that its marginal SHAP "
    "contribution is absorbed by the correlated features."
)

figure(
    os.path.join(IMG_DIR, "fig8_shap_importance.png"),
    "Figure XX. Global SHAP feature importance (mean |SHAP value|) for all "
    "three tree-based classifiers, colour-coded by marker category.",
    width_cm=16,
)

body(
    "Standardized Structure consistently ranks among the top two markers for "
    "all three tree-based classifiers in the SHAP analysis. Linguistic "
    "Homogeneity appears among the top two for all three models as well, "
    "ranking first for LightGBM (mean |SHAP| = 5.522). The absolute SHAP "
    "magnitudes differ substantially across classifiers: LightGBM produces "
    "the largest values (up to 5.522), reflecting its deeper-boosted tree "
    "structure, while RF operates on a much smaller scale "
    "(maximum 0.180) due to its probability-based output aggregation. "
    "Despite these scale differences, the relative feature rankings are "
    "broadly consistent, confirming that the same markers drive predictions "
    "regardless of the specific ensemble method."
)

# ── 4.4.3 SHAP beeswarm plots ───────────────────────────────────────────────

body(
    "Figures XX–XX present the SHAP beeswarm plots for XGBoost, RF, "
    "and LightGBM, respectively. Each dot represents a single review; the "
    "horizontal position indicates the SHAP value (impact on the model's "
    "log-odds output), and the colour encodes the marker score (red = high, "
    "blue = low). These plots reveal the direction of each marker's effect: "
    "high marker scores (red dots on the right) consistently push the model's "
    "prediction toward the AI class, while low marker scores (blue dots on "
    "the left) push toward the human class."
)

figure(
    os.path.join(IMG_DIR, "fig9_shap_beeswarm_xgboost.png"),
    "Figure XX. SHAP beeswarm plot, XGBoost. Each dot represents one review; "
    "horizontal position indicates the SHAP value, and colour encodes the "
    "feature value (red = high, blue = low).",
    width_cm=14,
)

figure(
    os.path.join(IMG_DIR, "fig9_shap_beeswarm_randomforest.png"),
    "Figure XX. SHAP beeswarm plot, RF.",
    width_cm=14,
)

figure(
    os.path.join(IMG_DIR, "fig9_shap_beeswarm_lightgbm.png"),
    "Figure XX. SHAP beeswarm plot, LightGBM.",
    width_cm=14,
)

body(
    "A consistent pattern across all three beeswarm plots is the strong "
    "bidirectional effect of Standardized Structure: reviews with high "
    "structural regularity (red) receive large positive SHAP values, "
    "decisively pushing the prediction toward AI, while reviews with low "
    "structural regularity (blue) receive large negative SHAP values, "
    "anchoring the prediction to the human class. The XGBoost beeswarm "
    "(Figure XX) shows Standardized Structure SHAP values spanning from "
    "approximately −6 to +3, by far the widest spread of any marker. "
    "Linguistic Homogeneity and Generic Domain Language exhibit similar "
    "directional patterns, though with more moderate magnitudes. "
    "In contrast, the bottom-ranked markers (Predictable Criticism, "
    "Absence of Personal Signals) show tightly clustered dots near zero, "
    "confirming their negligible marginal SHAP contribution."
)

body(
    "A notable difference across models emerges in the RF beeswarm "
    "(Figure XX): the feature ordering differs, with Conceptual Feedback "
    "ranking third (above Repetition Patterns), and Absence of Personal "
    "Signals appearing in sixth position with a non-trivial spread. "
    "This suggests that the bagged ensemble captures complementary feature "
    "interactions that the boosted methods subsume into the top markers. "
    "The LightGBM beeswarm (Figure XX) exhibits the largest SHAP magnitudes "
    "overall, with Linguistic Homogeneity values extending from −10 to +10, "
    "reflecting the model's deeper tree structure and more extreme log-odds."
)

# ── 4.4.4 Cross-classifier feature importance comparison ─────────────────────

body_mixed([
    ("Cross-classifier comparison. ", True, False),
    ("Figure XX presents the normalised feature importance comparison across "
     "all four classifiers (including LR). The chart reveals "
     "that Standardized Structure achieves the maximum normalised importance "
     "(1.0) in every classifier when native importances are rescaled. "
     "A key divergence appears in the secondary markers: tree-based methods "
     "show steep drop-offs after the top two features, whereas Logistic "
     "Regression distributes importance more uniformly across all eight markers. "
     "In particular, Repetition Patterns receives a normalised score of ≈ 0.57 "
     "in LR but varies widely across tree-based models "
     "(0.02 for XGBoost vs. 0.77 for RF). This divergence reflects "
     "the linear model's inability to exploit complex interactions: it "
     "compensates by hedging across all available features, whereas tree "
     "ensembles concentrate importance on the most informative markers.", False, False),
])

figure(
    os.path.join(IMG_DIR, "fig10_feature_importance_comparison.png"),
    "Figure XX. Normalised feature importance comparison across all four "
    "classifiers. Values are scaled relative to the maximum importance "
    "within each model to enable cross-model comparison.",
    width_cm=16,
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4.5 Adversarial Robustness
# ═══════════════════════════════════════════════════════════════════════════════

if H is not None:
    add_heading2("4.5. Adversarial robustness evaluation")

    body(
        "A critical test of the framework's practical utility is its performance on "
        "the adversarial subset, reviews generated with explicit instructions to "
        "evade detection. To assess robustness, we evaluated all four classifiers "
        f"on a hard subset containing only adversarial AI reviews (n = {N_HARD // 2}) "
        f"and a matched sample of human reviews (n = {N_HARD // 2}). "
        "Table XX reports the results."
    )

    caption(f"Table XX. Classifier performance on the adversarial hard subset (n = {N_HARD})")

    # Build adversarial table from hard_subset_comparison.json
    hard_rows = []
    for name in NAMES_ORDER:
        h = H[name]
        hard_rows.append([
            NICE_NAMES[name],
            f"{h['acc']:.3f}",
            f"{h['auc']:.3f}",
            str(h['tp']),
            str(h['fp']),
            str(h['fn']),
            str(h['tn']),
        ])

    add_table(
        headers=["Classifier", "Accuracy", "AUC-ROC", "TP", "FP", "FN", "TN"],
        rows=hard_rows,
    )

    # Extract key values for narrative
    h_rf = H["RandomForest"]

    body(
        f"All four classifiers maintain accuracy at or above {min(H[n]['acc'] for n in NAMES_ORDER):.2f} "
        f"on the adversarial subset, with AUC-ROC values exceeding "
        f"{min(H[n]['auc'] for n in NAMES_ORDER):.2f}. RF achieves the "
        f"highest performance ({h_rf['acc']:.3f} accuracy, {h_rf['auc']:.3f} AUC-ROC), "
        f"correctly classifying all {h_rf['tp']} adversarial reviews as AI-generated "
        f"with zero false negatives. The remaining classifiers each produce a single "
        f"false negative, indicating that only one adversarial review out of "
        f"{N_HARD // 2} successfully evades detection."
    )

    body(
        "These results demonstrate that the eight-marker taxonomy captures "
        "linguistic properties that are difficult to suppress even when the LLM "
        "is explicitly prompted to mimic human writing style. While the adversarial "
        "prompts successfully reduced marker scores to near-human levels on several "
        "dimensions (e.g., Standardized Structure, Absence of Personal Signals), "
        "the residual signal across the full marker profile remains sufficient for "
        "highly accurate detection."
    )

    body(
        f"Notably, even LR achieves {H['LogisticRegression']['acc']:.2f} "
        f"accuracy on the adversarial subset, substantially higher than its "
        f"{lr_m['acc']*100:.2f}% accuracy on the full dataset. This counter-intuitive "
        f"result likely reflects the smaller, more balanced evaluation set "
        f"({N_HARD // 2} vs. {N_HARD // 2}) and the absence of the class imbalance "
        f"that degrades LR's precision on the full dataset."
    )

    body_mixed([
        ("Implications for real-world deployment. ", True, False),
        ("The near-perfect performance on adversarial reviews suggests that the "
         "detection framework remains robust against current evasion strategies. "
         "However, we note two important caveats. First, the adversarial prompts, "
         "while sophisticated, represent a single point in the space of possible "
         "evasion strategies; more advanced techniques (e.g., iterative refinement "
         "based on detected markers) could potentially degrade performance. Second, "
         f"the hard subset evaluation is limited by its small sample size (n = {N_HARD}), "
         "and results should be interpreted with appropriate caution.", False, False),
    ])


# ═══════════════════════════════════════════════════════════════════════════════
# 4.6 RAG Retrieval Evaluation
# ═══════════════════════════════════════════════════════════════════════════════

if RAG is not None:
    add_heading2("4.6. RAG retrieval evaluation")

    # ── Compute RAG metrics from rag_evaluation.json ───────────────────────────
    human_qs = RAG["human_queries"]
    ai_qs = RAG["ai_queries"]
    n_human_q = len(human_qs)
    n_ai_q = len(ai_qs)

    # Top-1 accuracy: most_similar_label matches query class
    human_top1_correct = sum(1 for q in human_qs if q["most_similar_label"] == "Human")
    ai_top1_correct = sum(1 for q in ai_qs if q["most_similar_label"] == "AI-Generated")
    human_top1_acc = human_top1_correct / n_human_q * 100
    ai_top1_acc = ai_top1_correct / n_ai_q * 100
    overall_top1_acc = (human_top1_correct + ai_top1_correct) / (n_human_q + n_ai_q) * 100

    # Average top-k composition
    human_avg_human_in_topk = sum(q["human_matches"] for q in human_qs) / n_human_q
    human_avg_ai_in_topk = sum(q["ai_matches"] for q in human_qs) / n_human_q
    ai_avg_human_in_topk = sum(q["human_matches"] for q in ai_qs) / n_ai_q
    ai_avg_ai_in_topk = sum(q["ai_matches"] for q in ai_qs) / n_ai_q

    # Similarity scores
    human_sims = [q["avg_similarity"] for q in human_qs]
    ai_sims = [q["avg_similarity"] for q in ai_qs]
    human_mean_sim = sum(human_sims) / len(human_sims)
    ai_mean_sim = sum(ai_sims) / len(ai_sims)

    # Top-1 similarity scores
    human_top1_sims = [q["most_similar_score"] for q in human_qs]
    ai_top1_sims = [q["most_similar_score"] for q in ai_qs]
    human_mean_top1 = sum(human_top1_sims) / len(human_top1_sims)
    ai_mean_top1 = sum(ai_top1_sims) / len(ai_top1_sims)

    body(
        "The final stage of the framework employs Retrieval-Augmented Generation "
        "(RAG) to provide evidence-based support for the classifier's predictions. "
        "For each query review, the RAG module retrieves the most similar reviews "
        "from the knowledge base using FAISS with sentence-transformer embeddings "
        "(all-MiniLM-L6-v2). The retrieved reviews serve as contextual evidence: "
        "if a review classified as AI-generated is most similar to other AI-generated "
        "reviews in the knowledge base, this corroborates the classifier's decision "
        "and provides interpretable examples for the editor."
    )

    body(
        f"To evaluate the RAG component, we sampled {n_human_q} human and "
        f"{n_ai_q} AI-generated reviews as queries and retrieved the top-5 most "
        f"similar reviews from the full knowledge base (excluding the query itself). "
        f"Table XX reports the key retrieval metrics."
    )

    caption("Table XX. RAG retrieval evaluation metrics")

    add_table(
        headers=["Metric", "Human Queries\n(n = 100)", "AI Queries\n(n = 100)"],
        rows=[
            ["Top-1 accuracy", f"{human_top1_acc:.0f}%", f"{ai_top1_acc:.0f}%"],
            ["Avg same-class in top-5", f"{human_avg_human_in_topk:.1f}", f"{ai_avg_ai_in_topk:.1f}"],
            ["Avg cross-class in top-5", f"{human_avg_ai_in_topk:.1f}", f"{ai_avg_human_in_topk:.1f}"],
            ["Mean top-5 similarity", f"{human_mean_sim:.3f}", f"{ai_mean_sim:.3f}"],
            ["Mean top-1 similarity", f"{human_mean_top1:.3f}", f"{ai_mean_top1:.3f}"],
        ],
    )

    body(
        f"The top-1 retrieval accuracy, defined as the proportion of queries "
        f"whose nearest neighbour shares the same class label, reaches "
        f"{human_top1_acc:.0f}% for human queries and {ai_top1_acc:.0f}% for "
        f"AI-generated queries, yielding an overall accuracy of "
        f"{overall_top1_acc:.1f}%. This indicates that the semantic embedding "
        f"space naturally clusters reviews by origin, even though the embeddings "
        f"are computed from raw review text without access to the marker scores "
        f"or classifier labels."
    )

    body(
        f"The top-5 composition analysis further confirms this clustering effect. "
        f"For human queries, an average of {human_avg_human_in_topk:.1f} out of the "
        f"top-5 retrieved reviews are human-authored, with only "
        f"{human_avg_ai_in_topk:.1f} AI-generated reviews appearing in the "
        f"neighbourhood. For AI queries, the pattern mirrors: an average of "
        f"{ai_avg_ai_in_topk:.1f} of the top-5 neighbours are AI-generated, with "
        f"{ai_avg_human_in_topk:.1f} human reviews. This asymmetry in neighbourhood "
        f"composition provides a secondary signal that reinforces the classifier's "
        f"binary prediction."
    )

    body(
        f"The mean cosine similarity scores are comparable across query types: "
        f"{human_mean_sim:.3f} for human queries and {ai_mean_sim:.3f} for "
        f"AI-generated queries. This suggests that both classes exhibit similar "
        f"levels of intra-class semantic cohesion when measured through "
        f"general-purpose sentence embeddings. The slightly higher similarity "
        f"for human queries may reflect the larger pool of human reviews in the "
        f"knowledge base (5,772 vs. 2,000), which increases the likelihood of "
        f"finding close semantic matches."
    )

    figure(
        os.path.join(IMG_DIR, "fig13_rag_evaluation.png"),
        "Figure XX. RAG retrieval evaluation: (a) top-1 retrieval label "
        "accuracy for human and AI queries, (b) average top-5 neighbourhood "
        "composition by class, and (c) distribution of average cosine "
        "similarity scores across queries.",
        width_cm=16,
    )

    body_mixed([
        ("Role in the detection pipeline. ", True, False),
        ("The RAG module does not contribute to the binary classification "
         "decision itself; rather, it serves an interpretability and evidence "
         "function. When the classifier flags a review as potentially "
         "AI-generated, the retrieved similar reviews provide concrete examples "
         "that an editor can examine to understand why the review was flagged. "
         "The high top-1 accuracy demonstrates that the retrieved evidence "
         "is typically class-consistent, meaning the examples shown to the "
         "editor are genuinely representative of the predicted class. "
         "In the small number of cases where the nearest neighbour is from "
         "the opposite class, this may itself be informative, suggesting "
         "a boundary case that warrants closer editorial scrutiny.", False, False),
    ])


# ═══════════════════════════════════════════════════════════════════════════════
# Appendix: Prompt Design (for inclusion in the Methodology section)
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()

add_heading1("Appendix: Prompt Design")

body(
    "This appendix documents the two core prompts employed in the framework: "
    "the feature extraction prompt used to score the eight linguistic markers, "
    "and the adversarial generation prompt used to produce human-mimicking "
    "AI reviews. Both prompts represent critical methodological choices whose "
    "design directly affects the validity of the experimental results."
)

# ── A. Feature Extraction Prompt ────────────────────────────────────────────

add_heading2("A. Feature Extraction Prompt")

body(
    "A critical design choice in the LLM-as-feature-extractor paradigm is "
    "the framing of the extraction prompt. Early experiments revealed that "
    'prompts framing the LLM as "an expert in detecting AI-generated reviews" '
    "produced highly polarised scores — near-ceiling values (0.85–0.95) for "
    "AI-generated text and near-floor values (0.10–0.20) for human text — "
    "suggesting a self-recognition bias in which the LLM trivially identifies "
    "its own outputs. To mitigate this effect, we adopted a neutral framing "
    "that instructs the model to act as a linguistic analyst with no knowledge "
    "of the detection task."
)

body(
    "Each property is accompanied by a brief operational definition. For "
    "example, standardized_structure is defined as "
    '"How rigidly does the text follow a templated structure with clearly '
    'labeled sections (e.g., Summary, Strengths, Weaknesses)?", while '
    "absence_personal_signals is defined as "
    '"How absent are personal voice markers (e.g., \'I think\', \'I found\', '
    "'in my experience', expressions of uncertainty)?\". The model is "
    "instructed to return only a JSON object containing the eight numerical "
    "scores, with no justifications or additional text, to minimise parsing "
    "failures and ensure consistency across the corpus. The full prompt "
    "template is provided below."
)

caption("Prompt A. Feature extraction prompt template")

# Add the prompt as a code-style block
p_prompt = doc.add_paragraph()
p_prompt.paragraph_format.space_before = Pt(4)
p_prompt.paragraph_format.space_after = Pt(4)
p_prompt.paragraph_format.line_spacing = 1.0

EXTRACTION_PROMPT_TEXT = (
    'You are a linguistic analyst evaluating the writing characteristics '
    'of an academic peer review.\n\n'
    'Score each of the following 8 textual properties from 0.0 (not present '
    'at all) to 1.0 (very strongly present). Be precise and use the full '
    'range of scores.\n\n'
    'PROPERTIES:\n'
    '1. standardized_structure: How rigidly does the text follow a templated '
    'structure with clearly labeled sections (e.g., Summary, Strengths, '
    'Weaknesses)?\n'
    '2. predictable_criticism: How much does the text rely on common, '
    'formulaic critique phrases (e.g., "needs ablation study", "stronger '
    'baselines") rather than paper-specific criticism?\n'
    '3. excessive_balance: How diplomatically balanced is the tone? Does it '
    'systematically pair criticism with positive framing?\n'
    '4. linguistic_homogeneity: How uniform are the grammar, sentence length, '
    'and tone throughout the text?\n'
    '5. generic_domain_language: How much does the text use broad academic '
    'phrases (e.g., "novel approach", "significant contribution") rather '
    'than precise technical language?\n'
    '6. conceptual_feedback: How much does the feedback stay at a '
    'high/conceptual level without referencing specific lines, pages, '
    'figures, or tables?\n'
    '7. absence_personal_signals: How absent are personal voice markers '
    '(e.g., "I think", "I found", "in my experience", expressions of '
    'uncertainty)?\n'
    '8. repetition_patterns: How much repetitive or templated phrasing '
    'appears across sections?\n\n'
    'PEER REVIEW TEXT:\n'
    '"""\n'
    '{review_text}\n'
    '"""\n\n'
    'Respond ONLY with valid JSON containing the 8 scores (no justifications):\n'
    '{\n'
    '  "standardized_structure": 0.0,\n'
    '  "predictable_criticism": 0.0,\n'
    '  "excessive_balance": 0.0,\n'
    '  "linguistic_homogeneity": 0.0,\n'
    '  "generic_domain_language": 0.0,\n'
    '  "conceptual_feedback": 0.0,\n'
    '  "absence_personal_signals": 0.0,\n'
    '  "repetition_patterns": 0.0\n'
    '}'
)

run_p = p_prompt.add_run(EXTRACTION_PROMPT_TEXT)
run_p.font.name = "Courier New"
run_p.font.size = Pt(8)
run_p.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ── B. Adversarial Generation Prompt ────────────────────────────────────────

add_heading2("B. Adversarial Review Generation Prompt")

body(
    "To evaluate the robustness of the detection framework against deliberate "
    "evasion, we generated 1,000 adversarial reviews (500 via Claude Sonnet 4.6, "
    "500 via GPT 5.4) using a prompt specifically designed to suppress the "
    "linguistic markers targeted by the taxonomy. The adversarial prompt "
    "instructs the LLM to adopt the persona of an experienced human peer "
    "reviewer and provides explicit stylistic directives that directly "
    "counteract each marker category."
)

body(
    "The prompt specifies eleven style requirements, including: "
    "(i) avoiding markdown headers and templated structure (targeting "
    "standardized_structure), "
    '(ii) including personal voice markers such as "I think" and '
    '"after reading this twice" (targeting absence_personal_signals), '
    "(iii) referencing specific page numbers, figure numbers, and equation "
    "numbers (targeting conceptual_feedback), "
    "(iv) mixing formal and informal language with varied sentence lengths "
    "(targeting linguistic_homogeneity), "
    "(v) expressing genuine uncertainty (targeting excessive_balance), "
    "(vi) avoiding balanced Strengths/Weaknesses lists (targeting "
    "predictable_criticism), and "
    "(vii) including minor imperfections such as tangents and self-corrections "
    "(targeting repetition_patterns). "
    "Each review was generated based on the title, abstract, and content of "
    "a real ICLR submission drawn from the Gen-Review corpus, with temperature "
    "set to 0.9 to maximise stylistic diversity. The identical prompt was used "
    "for both Claude Sonnet 4.6 and GPT 5.4, enabling a controlled comparison "
    "of evasion effectiveness across LLM architectures."
)

caption("Prompt B. Adversarial review generation prompt template")

p_adv = doc.add_paragraph()
p_adv.paragraph_format.space_before = Pt(4)
p_adv.paragraph_format.space_after = Pt(4)
p_adv.paragraph_format.line_spacing = 1.0

ADVERSARIAL_PROMPT_TEXT = (
    'You are a real human peer reviewer at a top ML conference. You have been '
    'reviewing papers for 10+ years. Write a genuine, natural peer review for '
    'this paper.\n\n'
    'IMPORTANT STYLE REQUIREMENTS — your review must feel authentically human:\n'
    '- Do NOT use markdown headers like "### Summary" or "### Strengths". '
    'Write in flowing paragraphs or use simple formatting.\n'
    '- Include personal voice: "I think", "I found this confusing", '
    '"after reading this twice", "in my experience"\n'
    '- Reference specific parts: mention page numbers, figure numbers, '
    'equation numbers, line numbers, table numbers\n'
    '- Be inconsistent in tone — mix formal and informal language naturally\n'
    '- Show genuine uncertainty: "I might be wrong but...", '
    '"I\'m not entirely sure about..."\n'
    '- Vary your sentence length — mix short punchy sentences with longer ones\n'
    '- Include minor imperfections: slight tangents, self-corrections, '
    'strong opinions\n'
    '- Do NOT use balanced "Strengths/Weaknesses" lists. Real reviewers '
    'often focus more on one side\n'
    '- Occasionally be blunt or even a bit harsh — real reviewers sometimes are\n'
    '- Reference your own expertise or related work you\'ve read\n\n'
    'Title: {title}\n\n'
    'Abstract: {abstract}\n\n'
    'Content: {content}\n\n'
    'Write your review now. Make it 200-500 words, as a real busy reviewer would.'
)

run_a = p_adv.add_run(ADVERSARIAL_PROMPT_TEXT)
run_a.font.name = "Courier New"
run_a.font.size = Pt(8)
run_a.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

body(
    "As shown in Table XX (§4.2), both models produced remarkably similar "
    "marker profiles under this prompt, suggesting that the evasion strategies "
    "generalise across LLM architectures rather than being model-specific. "
    "The adversarial reviews successfully reduced marker scores to near-human "
    "levels on several dimensions (e.g., Standardized Structure: 0.077 for "
    "adversarial Claude vs. 0.205 for human), confirming that the prompt "
    "effectively induced human-mimicking behaviour."
)


# ── Save ─────────────────────────────────────────────────────────────────────
out_path = os.path.join(OUT_DIR, "research_results.docx")
doc.save(out_path)
print(f"✅ Saved to: {out_path}")
